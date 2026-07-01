from __future__ import annotations

import hashlib
import importlib.util
import json
import re
from dataclasses import dataclass
from importlib import metadata as importlib_metadata
from pathlib import Path
from typing import Any

from app.models.schemas import DocumentMetadata, ParsedArtifact, ParsedBlock, ParsedDocument, ParsedPage, SUPPORTED_EXTENSIONS


MANUAL_TYPES = ("AMM", "FIM", "TSM", "IPC", "WDM", "SRM", "CMM", "SB", "AD", "MEL", "CDL")
CHINESE_MANUAL_TYPES = {
    "维护手册": "维护手册",
    "维修手册": "维修手册",
    "操作手册": "操作手册",
    "飞行手册": "飞行手册",
    "排故手册": "排故手册",
}
DOCLING_FORMATS = {".pdf", ".docx"}
DOCLING_PAGE_BREAK = "\n\n<!-- aviation-kb-page-break -->\n\n"
DOCLING_RUNTIME_MODULES = ("docling", "docling_parse", "pypdfium2")
DOCLING_DOCUMENT_TIMEOUT_SECONDS: int | None = None


@dataclass(frozen=True)
class _ParseResult:
    pages: list[ParsedPage]
    artifact: ParsedArtifact
    blocks: list[ParsedBlock]


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def discover_files(path: Path) -> list[Path]:
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_EXTENSIONS else []
    files: list[Path] = []
    for child in path.rglob("*"):
        if child.is_file() and child.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(child)
    return sorted(files)


def parse_document(path: Path, overrides: DocumentMetadata | None = None) -> ParsedDocument:
    path = path.expanduser().resolve()
    suffix = path.suffix.lower()
    result = _parse_with_docling(path) if suffix in DOCLING_FORMATS else None
    if result is None:
        if suffix == ".pdf":
            pages = _parse_pdf(path)
            result = _fallback_parse_result(path, pages, parser_name="pypdf")
        elif suffix == ".docx":
            pages = _parse_docx(path)
            result = _fallback_parse_result(path, pages, parser_name="python-docx")
        elif suffix in {".txt", ".md", ".markdown"}:
            pages = _parse_text(path)
            result = _fallback_parse_result(path, pages, parser_name="plain-text")
        else:
            raise ValueError(f"Unsupported document type: {path.suffix}")

    pages = result.pages
    inferred = infer_metadata(path, "\n".join(page.text[:2000] for page in pages[:3]))
    metadata = overrides.merged_with(inferred) if overrides else inferred
    return ParsedDocument(path=path, pages=pages, metadata=metadata, artifact=result.artifact, blocks=result.blocks)


def docling_status() -> dict[str, object]:
    modules = {name: importlib.util.find_spec(name) is not None for name in DOCLING_RUNTIME_MODULES}
    active = all(modules.values())
    if active:
        note = "Docling will parse PDF/DOCX inputs before chunking."
    elif modules["docling"]:
        note = "Docling package is present, but PDF/DOCX backend modules are incomplete; using fallback parsers."
    else:
        note = "Docling is not installed in this environment; using fallback parsers."

    return {
        "preferred": "docling",
        "active": active,
        "fallback": "pypdf/python-docx",
        "formats": sorted(DOCLING_FORMATS),
        "modules": modules,
        "note": note,
    }


def infer_metadata(path: Path, sample: str = "") -> DocumentMetadata:
    raw_haystack = f"{path.name}\n{sample}"
    haystack = raw_haystack.upper()
    manual_type = next((kind for kind in MANUAL_TYPES if re.search(rf"\b{kind}\b", haystack)), None)
    if not manual_type:
        manual_type = next((value for key, value in CHINESE_MANUAL_TYPES.items() if key in raw_haystack), None)

    ata_match = re.search(r"第\s*(\d{2})\s*章", raw_haystack)
    if not ata_match:
        ata_match = re.search(r"\bATA[\s_-]*(\d{2})\b", haystack)
    if not ata_match:
        ata_match = re.search(r"(?:^|[^\d])(\d{2})[-_ ](?:\d{2}|[A-Z])", path.stem.upper())

    aircraft_match = re.search(
        r"(?<![A-Z0-9])"
        r"(CH[-_ ]?\d+[A-Z]?|A3\d{2}|A220|A350|A380|B7\d{2}|737NG|737MAX|ARJ21|C919|E19\d|E17\d)"
        r"(?![A-Z0-9])",
        haystack,
    )
    revision_match = re.search(r"\b(V\d+(?:\.\d+){0,3})\b", haystack)
    date_match = re.search(r"(\d{4})[-./年](\d{1,2})[-./月](\d{1,2})", raw_haystack)
    if not date_match:
        date_match = re.search(r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)", raw_haystack)

    contains_cjk = bool(re.search(r"[\u3400-\u9fff]", sample))
    contains_ascii = bool(re.search(r"[A-Za-z]", sample))
    language = "mixed" if contains_cjk and contains_ascii else "zh" if contains_cjk else "en" if contains_ascii else None

    return DocumentMetadata(
        manual_type=manual_type,
        aircraft_model=aircraft_match.group(1).replace("_", "-").replace(" ", "-") if aircraft_match else None,
        ata_chapter=ata_match.group(1) if ata_match else None,
        manual_revision=revision_match.group(1) if revision_match else None,
        effective_date=_normalized_date(date_match) if date_match else None,
        language=language,
    )


def _normalized_date(match: re.Match[str]) -> str:
    year, month, day = match.groups()
    return f"{int(year):04d}-{int(month):02d}-{int(day):02d}"


def _parse_pdf(path: Path) -> list[ParsedPage]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF parsing requires pypdf. Run `pip install pypdf`.") from exc

    reader = PdfReader(str(path))
    pages: list[ParsedPage] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(ParsedPage(page_no=index, text=text))
    return pages


def _parse_with_docling(path: Path) -> _ParseResult | None:
    if not _docling_backend_available():
        return None
    expected_pages = _pdf_page_count(path) if path.suffix.lower() == ".pdf" else 0
    try:
        converter = _build_docling_converter(do_ocr=False)
    except Exception:
        return None

    try:
        result = _convert_with_docling(converter, path, ocr_enabled=False)
        if _is_complete_docling_parse(result.pages, expected_pages) and _has_enough_docling_text(result.pages):
            return result

        result = _convert_with_docling(_build_docling_converter(do_ocr=True), path, ocr_enabled=True)
        return result if _is_complete_docling_parse(result.pages, expected_pages) else None
    except Exception:
        return None


def _build_docling_converter(do_ocr: bool):
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = do_ocr
    pipeline_options.document_timeout = DOCLING_DOCUMENT_TIMEOUT_SECONDS
    pipeline_options.force_backend_text = not do_ocr
    return DocumentConverter(format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)})


def _convert_with_docling(converter: object, path: Path, ocr_enabled: bool) -> _ParseResult:
    result = converter.convert(path)
    document = result.document
    markdown = document.export_to_markdown(
        page_break_placeholder=DOCLING_PAGE_BREAK,
        traverse_pictures=True,
        compact_tables=True,
    )
    raw_json = _docling_export_to_dict(document)
    pages = _pages_from_docling_document(document, markdown)
    blocks = _blocks_from_docling_document(document)
    if not blocks:
        blocks = _blocks_from_pages(pages)
    artifact = ParsedArtifact(
        parser_name="docling",
        parser_version=_module_version("docling"),
        source_format=path.suffix.lower().lstrip("."),
        status="succeeded",
        page_count=len(pages),
        markdown=_normalize_docling_markdown(markdown),
        raw_json=raw_json,
        ocr_enabled=ocr_enabled,
    )
    return _ParseResult(pages=pages, artifact=artifact, blocks=blocks)


def _pages_from_docling_document(document: object, markdown: str) -> list[ParsedPage]:
    pages = []
    if total_pages := _docling_page_count(document):
        for page_no in range(1, total_pages + 1):
            text = document.export_to_markdown(
                page_no=page_no,
                page_break_placeholder=DOCLING_PAGE_BREAK,
                traverse_pictures=True,
                compact_tables=True,
            )
            if text.strip():
                pages.append(ParsedPage(page_no=page_no, text=_normalize_docling_markdown(text)))
        if pages:
            return pages
    return _pages_from_docling_markdown(markdown)


def _has_enough_docling_text(pages: list[ParsedPage]) -> bool:
    return sum(len(page.text.strip()) for page in pages[:3]) >= 40


def _docling_backend_available() -> bool:
    return bool(docling_status()["active"])


def _docling_page_count(document: object) -> int:
    num_pages = getattr(document, "num_pages", None)
    if callable(num_pages):
        try:
            return int(num_pages())
        except (TypeError, ValueError):
            pass
    pages = getattr(document, "pages", None)
    if pages is not None:
        try:
            return len(pages)
        except TypeError:
            pass
    return 0


def _pdf_page_count(path: Path) -> int:
    try:
        from pypdf import PdfReader
        return len(PdfReader(str(path)).pages)
    except Exception:
        return 0


def _is_complete_docling_parse(pages: list[ParsedPage], expected_pages: int) -> bool:
    if expected_pages <= 0:
        return True
    last_page = max((page.page_no for page in pages), default=0)
    return last_page >= max(1, int(expected_pages * 0.9))


def _pages_from_docling_markdown(markdown: str) -> list[ParsedPage]:
    parts = [part.strip() for part in markdown.split(DOCLING_PAGE_BREAK) if part.strip()]
    if parts:
        return [ParsedPage(page_no=index, text=_normalize_docling_markdown(part)) for index, part in enumerate(parts, start=1)]
    return _pages_from_text(_normalize_docling_markdown(markdown), chars_per_page=3500)


def _docling_export_to_dict(document: object) -> dict[str, Any] | None:
    export = getattr(document, "export_to_dict", None)
    if not callable(export):
        return None
    try:
        data = export()
    except Exception:
        return None
    return _json_safe(data)


def _blocks_from_docling_document(document: object) -> list[ParsedBlock]:
    iterate = getattr(document, "iterate_items", None)
    if not callable(iterate):
        return []

    try:
        items = iterate(with_groups=True, traverse_pictures=True)
    except TypeError:
        items = iterate()

    blocks: list[ParsedBlock] = []
    section_stack: dict[int, str] = {}
    for item, level in items:
        raw = _json_safe(_model_dump(item))
        block_type = _docling_block_type(item, raw)
        text = _docling_item_text(item, block_type, document)
        if not text:
            continue

        if block_type in {"title", "section_header"}:
            _update_docling_section_stack(section_stack, max(1, int(level or 1)), text)

        table_json = _json_dumps(_docling_table_payload(item, raw)) if block_type == "table" else None
        blocks.append(
            ParsedBlock(
                block_index=len(blocks),
                page_no=_docling_item_page_no(item, raw),
                block_type=block_type,
                text=text,
                section_path=_docling_section_path(section_stack),
                bbox_json=_json_dumps(_docling_item_bbox(item, raw)),
                table_json=table_json,
                raw_json=_json_dumps(raw),
            )
        )
    return blocks


def _blocks_from_pages(pages: list[ParsedPage]) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    for page in pages:
        for paragraph in _page_paragraphs(page.text):
            blocks.append(
                ParsedBlock(
                    block_index=len(blocks),
                    page_no=page.page_no,
                    block_type=_infer_block_type(paragraph),
                    text=paragraph,
                )
            )
    return blocks


def _fallback_parse_result(path: Path, pages: list[ParsedPage], parser_name: str) -> _ParseResult:
    markdown = "\n\n".join(f"<!-- page {page.page_no} -->\n\n{page.text}" for page in pages)
    artifact = ParsedArtifact(
        parser_name=parser_name,
        parser_version=_module_version(parser_name),
        source_format=path.suffix.lower().lstrip("."),
        status="succeeded",
        page_count=len(pages),
        markdown=markdown,
        raw_json={
            "parser": parser_name,
            "source": str(path),
            "pages": [{"page_no": page.page_no, "chars": len(page.text)} for page in pages],
        },
    )
    return _ParseResult(pages=pages, artifact=artifact, blocks=_blocks_from_pages(pages))


def _normalize_docling_markdown(text: str) -> str:
    text = text.replace("\r\n", "\n")
    text = text.replace(DOCLING_PAGE_BREAK.strip(), " ")
    text = re.sub(r"<!--\s*aviation-kb-page-break\s*-->", " ", text)
    text = re.sub(r"<!--\s*image\s*-->", " ", text)
    text = re.sub(r"\\([_*#`|\\])", r"\1", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text.strip()


def _docling_item_text(item: object, block_type: str, document: object | None = None) -> str:
    if block_type == "table":
        export_markdown = getattr(item, "export_to_markdown", None)
        if callable(export_markdown):
            try:
                text = export_markdown(doc=document)
                if text and str(text).strip():
                    return str(text).strip()
            except TypeError:
                try:
                    text = export_markdown()
                    if text and str(text).strip():
                        return str(text).strip()
                except Exception:
                    pass
            except Exception:
                pass
    for attr in ("text", "orig"):
        value = getattr(item, attr, None)
        if isinstance(value, str) and value.strip():
            return _normalize_docling_markdown(value)
    caption_text = getattr(item, "caption_text", None)
    if callable(caption_text):
        try:
            value = caption_text()
            if value and str(value).strip():
                return _normalize_docling_markdown(str(value))
        except Exception:
            pass
    return ""


def _docling_block_type(item: object, raw: Any) -> str:
    label = str(getattr(item, "label", "") or _nested_value(raw, "label") or "").lower()
    class_name = item.__class__.__name__.lower()
    if "table" in class_name or "table" in label:
        return "table"
    if "title" in label:
        return "title"
    if "section" in label or "heading" in label:
        return "section_header"
    if "list" in label:
        return "list"
    if "caption" in label:
        return "caption"
    if "picture" in class_name or "picture" in label or "image" in label:
        return "picture"
    return "paragraph"


def _docling_item_page_no(item: object, raw: Any) -> int | None:
    prov = getattr(item, "prov", None) or _nested_value(raw, "prov")
    if isinstance(prov, list) and prov:
        first = prov[0]
        page_no = getattr(first, "page_no", None) if not isinstance(first, dict) else first.get("page_no")
        try:
            return int(page_no) if page_no is not None else None
        except (TypeError, ValueError):
            return None
    return None


def _docling_item_bbox(item: object, raw: Any) -> Any:
    prov = getattr(item, "prov", None) or _nested_value(raw, "prov")
    if isinstance(prov, list) and prov:
        first = prov[0]
        bbox = getattr(first, "bbox", None) if not isinstance(first, dict) else first.get("bbox")
        return _json_safe(_model_dump(bbox))
    return None


def _docling_table_payload(item: object, raw: Any) -> Any:
    data = getattr(item, "data", None) or _nested_value(raw, "data")
    return _json_safe(_model_dump(data))


def _update_docling_section_stack(section_stack: dict[int, str], level: int, heading: str) -> None:
    for key in list(section_stack):
        if key >= level:
            section_stack.pop(key)
    section_stack[level] = heading


def _docling_section_path(section_stack: dict[int, str]) -> str | None:
    if not section_stack:
        return None
    return " > ".join(section_stack[key] for key in sorted(section_stack))


def _page_paragraphs(text: str) -> list[str]:
    rough = re.split(r"\n\s*\n", text.replace("\r\n", "\n"))
    return [item.strip() for item in rough if item.strip()]


def _infer_block_type(text: str) -> str:
    upper = text.upper()
    if "WARNING" in upper or "警告" in text:
        return "warning"
    if "CAUTION" in upper or "注意" in text:
        return "caution"
    if "NOTE" in upper or "注：" in text or "注:" in text:
        return "note"
    if "|" in text and text.count("|") >= 2:
        return "table"
    return "paragraph"


def _model_dump(value: Any) -> Any:
    if value is None:
        return None
    dump = getattr(value, "model_dump", None)
    if callable(dump):
        try:
            return dump(mode="json")
        except TypeError:
            return dump()
        except Exception:
            return str(value)
    return value


def _json_safe(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(key): _json_safe(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [_json_safe(item) for item in value]
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    return _json_safe(_model_dump(value)) if hasattr(value, "model_dump") else str(value)


def _json_dumps(value: Any) -> str | None:
    if value in (None, "", [], {}):
        return None
    return json.dumps(value, ensure_ascii=False, sort_keys=True)


def _nested_value(value: Any, key: str) -> Any:
    if isinstance(value, dict):
        return value.get(key)
    return None


def _module_version(module_name: str) -> str | None:
    package_names = {
        "python-docx": "python-docx",
        "plain-text": None,
    }
    package_name = package_names.get(module_name, module_name)
    if package_name is None:
        return None
    try:
        return importlib_metadata.version(package_name)
    except importlib_metadata.PackageNotFoundError:
        return None


def _parse_docx(path: Path) -> list[ParsedPage]:
    try:
        from docx import Document
    except ImportError:
        return _parse_docx_with_textutil(path)

    document = Document(str(path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    return _pages_from_text("\n".join(blocks), chars_per_page=3500)


def _parse_docx_with_textutil(path: Path) -> list[ParsedPage]:
    import subprocess
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    return _pages_from_text(result.stdout, chars_per_page=3500)


def _parse_text(path: Path) -> list[ParsedPage]:
    text = path.read_text(encoding="utf-8", errors="replace")
    if "\f" in text:
        return [ParsedPage(page_no=index, text=page) for index, page in enumerate(text.split("\f"), start=1)]
    return _pages_from_text(text, chars_per_page=3500)


def _pages_from_text(text: str, chars_per_page: int) -> list[ParsedPage]:
    text = text.strip()
    if not text:
        return [ParsedPage(page_no=1, text="")]
    pages = []
    for index, start in enumerate(range(0, len(text), chars_per_page), start=1):
        pages.append(ParsedPage(page_no=index, text=text[start : start + chars_per_page]))
    return pages
